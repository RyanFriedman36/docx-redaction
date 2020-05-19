
"""

File: redact.py

Project: Automatic legal-document redactionn desktop application

Author: Ryan Friedman

Description: This program will allow a user to supply a Microsoft Word
             Document and a list of proper nouns that need to be redacted.
             It will create a new Word Document with any version of the proper
             nouns replaced with a black highlight. We create a new document to
             ensure fresh metadata.

"""
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import tkinter as tk
from tkinter import filedialog
from string import punctuation
import os
from os import remove
import copy


def binary_search(source, target):
    """ This function searches for a target in source in logarithmic time """

    left = 0
    right = len(source) - 1

    while left <= right:
        median = left + ((right - left) // 2)
        if source[median] == target:
            return True
        elif source[median] < target:
            left = median + 1
        else:
            right = median - 1

    return False


def popupmsg(msg):
    """ This function opens a pop-up window and displays a message. The user
        may click 'okay' to close the window. """

    popup = tk.Tk()
    popup.wm_title("Auto-Redact")
    label = tk.Label(popup, text=msg, font=("Helvetica", 10))
    label.pack(side="top", fill="x", pady=10)
    B1 = tk.Button(popup, text="Okay", command=popup.destroy)
    B1.pack()
    popup.mainloop()


def requestFile():
    """ This function opens a file selection window using tkinter and
    returns the file path """

    root = tk.Tk()
    root.withdraw()

    file_path = filedialog.askopenfilename()

    if not file_path:
        exit(1)

    return file_path


def getDirFromFile(file_path):
    """ This function finds the last forward or back slash in directory path
        and slices the file name from it """

    break_index = -1
    for i in range(len(file_path) - 1, -1, -1):
        if file_path[i] == "\\" or file_path[i] == "/":
            break_index = i
            break

    if break_index != -1:
        file_path = file_path[:break_index + 1]

    return file_path


def processInfoFile(info_path):
    """ Opens the file and cleans/processes information to be redacted
        into the proper state """

    redact_file = open(info_path, "r")
    redact_info = redact_file.read()
    redact_file.close()

    redact_info = redact_info.split(" ")

    for i in range(len(redact_info)):
        redact_info[i] = redact_info[i].strip(",")

    for entry in redact_info:
        if entry == "" or entry == "\n":
            redact_info.remove(entry)
        if " " in entry:
            entry = entry.strip()

    redact_info.sort()

    return redact_info


def processFiles(file_path_1, file_path_2):
    """ file_path_1 is our Word document and file_path_2 is any text file
        with proper nouns listed on separate lines. This is where we create
        our redacted Word File """

    redact_info = processInfoFile(file_path_2)

    doc = Document(file_path_1)

    new_doc = Document()
    temp_doc = Document()

    for i in range(len(doc.paragraphs)):
        processPara(doc.paragraphs[i], redact_info, new_doc, temp_doc)

    new_doc.save(getDirFromFile(file_path_1) + "redacted version.docx")


def processPara(para, redact_info, new_doc, temp_doc):
    """ Gets indices for words that are to be redacts them then returns a
        a paragraph with those words redacted  """

    # gets indices within para.text of sensitive info
    redact_indices = getRedactIndices(para, redact_info)

    # creates a copy of para where each char is its own run
    para = convertRuns(para, temp_doc)

    redact(para, redact_indices, new_doc)


def getRedactIndices(para, redact_info):
    """ Returns a list containing tuples that represent the indices that any
        instance of redact_info span in para """

    redact_indices = []

    curr_word = ""
    for i in range(len(para.text)):
        char = para.text[i]
        if char == " " or char == "\t" or char == "\n":
            curr_word = curr_word.strip(punctuation)
            if binary_search(redact_info, curr_word):
                redact_indices.append((i - len(curr_word), i))
            curr_word = ""
        else:
            curr_word += char

    # check the final word in each paragraph
    if curr_word:
        curr_word = curr_word.strip(punctuation)
        if binary_search(redact_info, curr_word):
            redact_indices.append((i - len(curr_word) + 1, len(para.text)))

    return redact_indices


def convertRuns(para, temp_doc):
    """ Converts all text in a paragraph to individual runs for each character
        so that we may uniformly redact instances of sensitive information."""

    p = temp_doc.add_paragraph()

    for run in para.runs:
        for char in run.text:
            r = p.add_run(char)
            r.bold = run.bold
            r.italic = run.italic
            r.underline = run.underline
            r.font.color.rgb = run.font.color.rgb
            r.font.name = run.font.name
            r.style.name = run.style.name
            r.font.size = run.font.size

    p.paragraph_format.alignment = para.paragraph_format.alignment
    p.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
    p.paragraph_format.keep_together = para.paragraph_format.keep_together
    p.paragraph_format.keep_with_next = para.paragraph_format.keep_with_next
    p.paragraph_format.left_indent = para.paragraph_format.left_indent
    p.paragraph_format.line_spacing = para.paragraph_format.line_spacing
    p.paragraph_format.line_spacing_rule = para.paragraph_format.line_spacing_rule
    p.paragraph_format.page_break_before = para.paragraph_format.page_break_before
    p.paragraph_format.right_indent = para.paragraph_format.right_indent
    p.paragraph_format.space_after = para.paragraph_format.space_after
    p.paragraph_format.widow_control = para.paragraph_format.widow_control
    p.paragraph_format.space_before = para.paragraph_format.space_before

    return p


def redact(para, redact_indices, new_doc):
    """ This function modifies our paragraph object with new 'black runs'
        which are our redactions """

    index_map = getIndexMap(para, redact_indices)

    p = new_doc.add_paragraph()

    for i in range(len(para.runs)):
        if index_map[i] == 1:
            run = para.runs[i]
            r = p.add_run("X")
            r.font.highlight_color = WD_COLOR_INDEX.BLACK
            r.bold = run.bold
            r.italic = run.italic
            r.underline = run.underline
            r.font.color.rgb = run.font.color.rgb
            r.font.name = run.font.name
            r.style.name = run.style.name
            r.font.size = run.font.size
        else:
            run = para.runs[i]
            r = p.add_run(run.text)
            r.bold = run.bold
            r.italic = run.italic
            r.underline = run.underline
            r.font.color.rgb = run.font.color.rgb
            r.font.name = run.font.name
            r.style.name = run.style.name
            r.font.size = run.font.size

    p.paragraph_format.alignment = para.paragraph_format.alignment
    p.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
    p.paragraph_format.keep_together = para.paragraph_format.keep_together
    p.paragraph_format.keep_with_next = para.paragraph_format.keep_with_next
    p.paragraph_format.left_indent = para.paragraph_format.left_indent
    p.paragraph_format.line_spacing = para.paragraph_format.line_spacing
    p.paragraph_format.line_spacing_rule = para.paragraph_format.line_spacing_rule
    p.paragraph_format.page_break_before = para.paragraph_format.page_break_before
    p.paragraph_format.right_indent = para.paragraph_format.right_indent
    p.paragraph_format.space_after = para.paragraph_format.space_after
    p.paragraph_format.widow_control = para.paragraph_format.widow_control
    p.paragraph_format.space_before = para.paragraph_format.space_before



def getIndexMap(para, redact_indices):
    index_map = {}

    for i in range(len(para.text)):
        index_map[i] = 0

    for entry in redact_indices:
        for i in range(entry[0], entry[1]):
            index_map[i] = 1

    return index_map


def main():


    # popupmsg("Please select the Word document that you would like to redact from.")
    file_path_1 = requestFile()

    file_path_2 = requestFile()

    processFiles(file_path_1, file_path_2)



    # TESTING:

    # my_doc = Document()
    # for i in range(10):
    #     my_doc.add_paragraph(str(i))


    # my_doc.paragraphs[1] = my_doc.paragraphs[1].add_run("adhfalsdfj").bold = True
    # my_doc.paragraphs[1].runs[-1].font.highlight_color = WD_COLOR_INDEX.BLACK
    # my_doc.paragraphs[1].runs[-1].bold = False

    # for para in my_doc.paragraphs:
    #     print(para.text)

    # target_dir = getDirFromFile(requestFile())
    # os.remove(target_dir + "output_test.docx")

    # my_doc.save(target_dir + "output_test.docx")




    # popupmsg("The redacted version of your file has been created.")

    # display directory?

    # ask if they would like to open it in word

    # ask if they would like you to port to a pdf?


if __name__ == "__main__":
    main()